#!/usr/bin/env python3
"""
ZUWE-ACM — Flask 后端
提供题目 API 和代码评测 API。
"""

import os
import sys
import json
import re
import io
import time as _time
import traceback
import threading
import requests as http_requests

from flask import Flask, jsonify, request, render_template, Response
from docx import Document

# PyInstaller 打包后数据文件在 sys._MEIPASS 中
if getattr(sys, 'frozen', False):
    _ROOT = sys._MEIPASS
else:
    _ROOT = os.path.dirname(os.path.abspath(__file__))

# 尝试从 .env 文件加载环境变量（如果存在）
_env_path = os.path.join(_ROOT, ".env")
if os.path.exists(_env_path):
    with open(_env_path, "r") as _f:
        for _line in _f:
            _line = _line.strip()
            if _line and not _line.startswith("#") and "=" in _line:
                _k, _v = _line.split("=", 1)
                os.environ.setdefault(_k.strip(), _v.strip())

# 添加 sandbox 到路径
sys.path.insert(0, _ROOT)
from sandbox.runner import run_with_test_cases
import sandbox.runner as _sandbox_runner

app = Flask(__name__)
app.config['MAX_CONTENT_LENGTH'] = 50 * 1024 * 1024  # 请求体上限 50MB

# ── CSRF 保护：只允许来自本机的 API 请求 ──
@app.before_request
def _check_origin():
    if request.method in ('GET', 'HEAD'):
        return
    origin = (request.headers.get('Origin', '') or '').strip()
    referer = (request.headers.get('Referer', '') or '').strip()
    allowed_prefixes = ['http://127.0.0.1:', 'http://localhost:', 'http://[::1]:']
    is_allowed = any(origin.startswith(p) or referer.startswith(p) for p in allowed_prefixes)
    if origin and not is_allowed:
        return jsonify({"error": "拒绝来自外部的请求"}), 403

# ── 心跳检测（关闭浏览器后自动退出程序） ──
_last_heartbeat = _time.time()
_quitting = False
_HEARTBEAT_TIMEOUT = 30  # 秒，给 AI 请求留足时间

def _heartbeat_watchdog():
    while True:
        _time.sleep(5)
        if _time.time() - _last_heartbeat > _HEARTBEAT_TIMEOUT:
            os._exit(0)

threading.Thread(target=_heartbeat_watchdog, daemon=True).start()

# 加载题目
QUESTIONS_PATH = os.path.join(_ROOT, "questions", "questions.json")
with open(QUESTIONS_PATH, "r", encoding="utf-8") as f:
    _builtin_questions = json.load(f)

# ── 用户数据目录（持久化所有改动） ──
def _get_user_data_dir():
    if sys.platform == "win32":
        base = os.environ.get("APPDATA", os.path.expanduser("~"))
        return os.path.join(base, "ZUWE-ACM")
    elif sys.platform == "darwin":
        return os.path.join(os.path.expanduser("~"), "Library", "Application Support", "ZUWE-ACM")
    else:
        return os.path.join(os.path.expanduser("~"), ".config", "ZUWE-ACM")

USER_DATA_DIR = _get_user_data_dir()
USER_QUESTIONS_PATH = os.path.join(USER_DATA_DIR, "questions.json")
SETTINGS_PATH = os.path.join(USER_DATA_DIR, "settings.json")

# 首次启动：复制内置题到用户数据目录
os.makedirs(USER_DATA_DIR, exist_ok=True)
if not os.path.exists(USER_QUESTIONS_PATH):
    with open(USER_QUESTIONS_PATH, "w", encoding="utf-8") as f:
        json.dump(_builtin_questions, f, ensure_ascii=False, indent=2)

# ── 用户设置（编译器路径等）──
def _load_settings():
    try:
        with open(SETTINGS_PATH, "r", encoding="utf-8") as f:
            return json.load(f)
    except (FileNotFoundError, json.JSONDecodeError):
        return {}

def _save_settings(data):
    try:
        with open(SETTINGS_PATH, "w", encoding="utf-8") as f:
            json.dump(data, f, ensure_ascii=False, indent=2)
    except OSError:
        pass

# ── 运行日志 ──
_LOG_DIR = os.path.join(USER_DATA_DIR, "log")
os.makedirs(_LOG_DIR, exist_ok=True)
_LOG_FILE = os.path.join(_LOG_DIR, "oj.log")

def _user_error(msg, detail=""):
    """返回用户友好的错误响应，同时记录详细日志。"""
    _log("ERROR", msg, detail)
    return jsonify({"error": msg, "detail": detail if detail else None}), 500


def _log(level, msg, *args):
    """简易日志：写入日志文件，不依赖 logging 模块。"""
    try:
        line = f"[{_time.strftime('%Y-%m-%d %H:%M:%S')}] [{level}] {msg}"
        if args:
            line += " | " + " | ".join(str(a) for a in args)
        with open(_LOG_FILE, "a", encoding="utf-8") as f:
            f.write(line + "\n")
    except OSError:
        pass

# 载入自定义编译器路径
_settings = _load_settings()
_custom_compiler = _settings.get("compiler_path", "").strip()
if _custom_compiler:
    _sandbox_runner.CPP_COMPILER = _custom_compiler
    _log("INFO", f"使用自定义编译器路径: {_custom_compiler}")


def _load_questions():
    """从用户数据目录加载题目（含所有增删改）。"""
    with open(USER_QUESTIONS_PATH, "r", encoding="utf-8") as f:
        return json.load(f)


def _save_questions(all_questions):
    """持久化所有题目到用户数据目录。"""
    with open(USER_QUESTIONS_PATH, "w", encoding="utf-8") as f:
        json.dump(all_questions, f, ensure_ascii=False, indent=2)


def _next_question_id():
    max_id = max((q["id"] for q in questions), default=0)
    return max_id + 1


# 加载题目到内存
questions = _load_questions()


@app.route("/")
def index():
    return render_template("index.html")


@app.route("/api/questions")
def get_questions():
    """返回题目列表（不含测试用例的答案）。"""
    summary = []
    for q in questions:
        summary.append({
            "id": q["id"],
            "title": q["title"],
            "difficulty": q["difficulty"],
            "category": q["category"],
        })
    return jsonify(summary)


@app.route("/api/questions/<int:qid>")
def get_question(qid):
    """返回单个题目详情（不含测试用例的答案和完整代码）。"""
    for q in questions:
        if q["id"] == qid:
            resp = {
                "id": q["id"],
                "title": q["title"],
                "difficulty": q["difficulty"],
                "category": q["category"],
                "description": q["description"],
                "input_format": q["input_format"],
                "output_format": q["output_format"],
                "sample_input": q["sample_input"],
                "sample_output": q["sample_output"],
                "hint": q["hint"],
                "template_code": q["template_code"],
                "template_code_cpp": q.get("template_code_cpp", ""),
                "test_cases": [
                    {"input": tc["input"]} for tc in q["test_cases"]
                ],
                "test_case_count": len(q["test_cases"]),
            }
            return jsonify(resp)
    return jsonify({"error": "题目不存在"}), 404


@app.route("/api/submit", methods=["POST"])
def submit():
    """提交代码并运行测试用例。"""
    data = request.get_json()
    qid = data.get("question_id")
    code = data.get("code", "")
    language = data.get("language", "python")

    if not qid or not code:
        return jsonify({"error": "缺少题目 ID 或代码"}), 400

    question = None
    for q in questions:
        if q["id"] == qid:
            question = q
            break

    if not question:
        return jsonify({"error": "题目不存在"}), 404

    test_cases = question["test_cases"]
    timeout = 5

    _log("INFO", f"评测请求", f"题目={qid}", f"语言={language}", f"测试数={len(test_cases)}")
    try:
        results = run_with_test_cases(code, test_cases, timeout, language)
        passed_count = sum(1 for r in results if r["passed"])
        total = len(results)
        _log("INFO", f"评测完成", f"题目={qid}", f"通过={passed_count}/{total}")

        return jsonify(
            {
                "passed_count": passed_count,
                "total": total,
                "accepted": passed_count == total,
                "results": results,
            }
        )
    except Exception as e:
        _log("ERROR", f"评测异常", f"题目={qid}", str(e))
        traceback.print_exc()
        err_msg = str(e)
        if "timed out" in err_msg.lower() or "timeout" in err_msg.lower():
            return jsonify({"error": "评测超时，请优化代码效率"}), 500
        if "memory" in err_msg.lower():
            return jsonify({"error": "内存超限，请检查是否有无限递归或大数组分配"}), 500
        if "compilation" in err_msg.lower() or "compile" in err_msg.lower():
            return jsonify({"error": "编译错误，请检查代码语法", "detail": err_msg}), 500
        return jsonify({"error": f"评测异常: {err_msg[:200]}"}), 500


@app.route("/api/submit-hidden", methods=["POST"])
def submit_hidden():
    """提交代码并运行隐藏测试用例（不返回详细结果）。"""
    data = request.get_json()
    qid = data.get("question_id")
    code = data.get("code", "")
    language = data.get("language", "python")

    if not qid or not code:
        return jsonify({"error": "缺少题目 ID 或代码"}), 400

    question = None
    for q in questions:
        if q["id"] == qid:
            question = q
            break

    if not question:
        return jsonify({"error": "题目不存在"}), 404

    test_cases = question.get("hidden_test_cases", question["test_cases"])
    if not test_cases:
        return jsonify({"error": "暂无隐藏测试用例"}), 400

    timeout = 5

    _log("INFO", f"隐藏评测请求", f"题目={qid}", f"语言={language}", f"测试数={len(test_cases)}")
    try:
        results = run_with_test_cases(code, test_cases, timeout, language)
        passed_count = sum(1 for r in results if r["passed"])
        total = len(results)
        _log("INFO", f"隐藏评测完成", f"题目={qid}", f"通过={passed_count}/{total}")

        return jsonify(
            {
                "passed_count": passed_count,
                "total": total,
                "accepted": passed_count == total,
                "results": [
                    {"test_id": r["test_id"], "passed": r["passed"], "time": r.get("time", 0)}
                    for r in results
                ],
            }
        )
    except Exception as e:
        _log("ERROR", f"隐藏评测异常", f"题目={qid}", str(e))
        traceback.print_exc()
        err_msg = str(e)
        if "timed out" in err_msg.lower() or "timeout" in err_msg.lower():
            return jsonify({"error": "评测超时，请优化代码效率"}), 500
        if "memory" in err_msg.lower():
            return jsonify({"error": "内存超限，请检查是否有无限递归或大数组分配"}), 500
        if "compilation" in err_msg.lower() or "compile" in err_msg.lower():
            return jsonify({"error": "编译错误，请检查代码语法", "detail": err_msg}), 500
        return jsonify({"error": f"评测异常: {err_msg[:200]}"}), 500


# ──────────────────────────────────────────
#  SSE 流式评测（逐条显示测试结果）
# ──────────────────────────────────────────

def _get_question_by_id(qid):
    for q in questions:
        if q["id"] == qid:
            return q
    return None


def _stream_test_events(code, test_cases, timeout, language, hide_details=False):
    """Generator: 逐条运行测试用例并通过 SSE 推送。"""
    total = len(test_cases)
    passed_count = 0
    start_ts = _time.time()
    error_msg = None

    yield f"data: {json.dumps({'type': 'start', 'total': total})}\n\n"

    try:
        for i, tc in enumerate(test_cases):
            results = run_with_test_cases(code, [tc], timeout, language)
            result = results[0]
            passed = result["passed"]

            if passed:
                passed_count += 1

            event = {
                "type": "result",
                "test_id": i + 1,
                "passed": passed,
                "time": result["time"],
            }
            if not hide_details:
                event["input"] = result.get("input", tc["input"].strip())
                event["expected"] = result.get("expected", tc["expected"].strip())
                event["actual"] = result.get("actual", "")
                event["error"] = result.get("error")

            yield f"data: {json.dumps(event, ensure_ascii=False)}\n\n"
    except Exception as e:
        traceback.print_exc()
        error_msg = str(e)
        _log("ERROR", "流式评测异常", error_msg)
        yield f"data: {json.dumps({'type': 'error', 'message': error_msg})}\n\n"

    elapsed_ms = int((_time.time() - start_ts) * 1000)
    yield f"data: {json.dumps({
        'type': 'done',
        'passed_count': passed_count,
        'total': total,
        'accepted': passed_count == total,
        'elapsed': elapsed_ms,
        'error': error_msg,
    }, ensure_ascii=False)}\n\n"


@app.route("/api/submit-stream", methods=["POST"])
def submit_stream():
    """SSE 流式评测可见测试用例。"""
    data = request.get_json()
    qid = data.get("question_id")
    code = data.get("code", "")
    language = data.get("language", "python")
    if not qid or not code:
        return jsonify({"error": "缺少题目 ID 或代码"}), 400
    question = _get_question_by_id(qid)
    if not question:
        return jsonify({"error": "题目不存在"}), 404
    def _gen():
        yield from _stream_test_events(code, question["test_cases"], 5, language, hide_details=False)
    return Response(_gen(), mimetype="text/event-stream",
                    headers={"Cache-Control": "no-cache", "X-Accel-Buffering": "no"})


@app.route("/api/submit-hidden-stream", methods=["POST"])
def submit_hidden_stream():
    """SSE 流式评测隐藏测试用例（提交后展示输入/输出详情）。"""
    data = request.get_json()
    qid = data.get("question_id")
    code = data.get("code", "")
    language = data.get("language", "python")
    if not qid or not code:
        return jsonify({"error": "缺少题目 ID 或代码"}), 400
    question = _get_question_by_id(qid)
    if not question:
        return jsonify({"error": "题目不存在"}), 404
    test_cases = question.get("hidden_test_cases", question["test_cases"])
    if not test_cases:
        return jsonify({"error": "暂无隐藏测试用例"}), 400
    def _gen():
        yield from _stream_test_events(code, test_cases, 5, language, hide_details=True)
    return Response(_gen(), mimetype="text/event-stream",
                    headers={"Cache-Control": "no-cache", "X-Accel-Buffering": "no"})


@app.route("/api/questions", methods=["POST"])
def create_question():
    """新增题目。请求体为标准题目 JSON（不含 id，自动分配）。"""
    data = request.get_json()
    required = ["title", "difficulty", "category", "description",
                "input_format", "output_format", "sample_input", "sample_output",
                "template_code", "test_cases"]
    for field in required:
        if field not in data:
            return jsonify({"error": f"缺少必填字段: {field}"}), 400

    if not isinstance(data["test_cases"], list) or len(data["test_cases"]) == 0:
        return jsonify({"error": "test_cases 必须为非空数组"}), 400

    for tc in data["test_cases"]:
        if "input" not in tc or "expected" not in tc:
            return jsonify({"error": "测试用例缺少 input 或 expected"}), 400

    global questions
    new_id = _next_question_id()
    new_question = {
        "id": new_id,
        "title": data["title"],
        "difficulty": data["difficulty"],
        "category": data["category"],
        "description": data["description"],
        "input_format": data["input_format"],
        "output_format": data["output_format"],
        "sample_input": data["sample_input"],
        "sample_output": data["sample_output"],
        "hint": data.get("hint", ""),
        "template_code": data["template_code"],
        "template_code_cpp": data.get("template_code_cpp", STD_TEMPLATE_CPP),
        "test_cases": data["test_cases"],
    }
    questions.append(new_question)
    _save_questions(questions)
    _log("INFO", f"新增题目", f"id={new_id}", f"标题={data['title']}")
    return jsonify({"message": f"题目 #{new_id} 已添加", "id": new_id}), 201


@app.route("/api/questions/<int:qid>", methods=["DELETE"])
def delete_question(qid):
    """删除指定题目并持久化到 questions.json。"""
    global questions
    for i, q in enumerate(questions):
        if q["id"] == qid:
            deleted = questions.pop(i)
            _save_questions(questions)
            _log("INFO", f"删除题目", f"id={qid}", f"标题={deleted['title']}")
            return jsonify({"message": f"已删除题目 #{qid}: {deleted['title']}"})
    return jsonify({"error": "题目不存在"}), 404


@app.route("/api/questions/batch-delete", methods=["POST"])
def batch_delete_questions():
    """批量删除题目。"""
    global questions
    data = request.get_json()
    ids = data.get("ids", [])
    if not ids:
        return jsonify({"error": "请指定要删除的题目 ID"}), 400
    ids = [int(i) for i in ids]
    deleted = []
    remaining = []
    for q in questions:
        if q["id"] in ids:
            deleted.append(q)
        else:
            remaining.append(q)
    if not deleted:
        return jsonify({"error": "未找到匹配的题目"}), 404
    questions = remaining
    _save_questions(questions)
    _log("INFO", f"批量删除", f"数量={len(deleted)}", f"ids={[d['id'] for d in deleted]}")
    return jsonify({"message": f"已删除 {len(deleted)} 道题目", "deleted_count": len(deleted)})


# ──────────────────────────────────────────
#  Word (.docx) 上传/下载
# ──────────────────────────────────────────

STD_TEMPLATE_PY = "def main():\n    # code here\n    pass\n\n\nif __name__ == '__main__':\n    main()"
STD_TEMPLATE_CPP = '#include<bits/stdc++.h>\n\nusing namespace std;\n\nint main( )\n{\n    return 0;\n}'

DOCX_LABELS = {
    "标题": "title",
    "难度": "difficulty",
    "分类": "category",
    "题目描述": "description",
    "输入格式": "input_format",
    "输出格式": "output_format",
    "样例输入": "sample_input",
    "样例输出": "sample_output",
    "提示": "hint",
}


def _parse_docx(fileobj) -> dict:
    """解析上传的 .docx，返回 question dict（不含 id）。"""
    doc = Document(fileobj)
    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]

    # 第一遍：提取单行标签字段
    data = {}
    for label, key in DOCX_LABELS.items():
        pattern = re.compile(re.escape(label) + r"[：:]\s*(.*)")
        for p in paragraphs:
            m = pattern.match(p)
            if m:
                data[key] = m.group(1).strip()
                break
        if key not in data:
            data[key] = "" if key == "hint" else None

    # 第二遍：用多行模式提取需要换行的字段
    multiline_keys = {
        "题目描述": "description",
        "输入格式": "input_format",
        "输出格式": "output_format",
        "样例输入": "sample_input",
        "样例输出": "sample_output",
    }
    for label, key in multiline_keys.items():
        content = _extract_multiline(paragraphs, label)
        if content:
            data[key] = content

    # 第三遍：解析测试用例
    test_cases = _parse_test_cases(paragraphs)

    missing = [k for k, v in data.items() if v is None]
    if missing:
        raise ValueError(f"缺少必填字段: {', '.join(missing)}")
    if not test_cases:
        raise ValueError("未找到测试用例")

    # 自动填入标准模版代码
    data.setdefault("template_code", STD_TEMPLATE_PY)
    data["test_cases"] = test_cases
    return data


def _extract_multiline(paragraphs: list, label: str) -> str | None:
    """提取标签后的多行内容（直到下一个标签或 '---' 为止）。"""
    pattern = re.compile(re.escape(label) + r"[：:]")
    start_idx = None
    for i, p in enumerate(paragraphs):
        if pattern.match(p):
            start_idx = i + 1
            break
    if start_idx is None:
        return None

    lines = []
    for p in paragraphs[start_idx:]:
        # 碰到其他标签或用例分隔符就停止
        if re.match(r"^(?:标题|难度|分类|题目描述|输入格式|输出格式|样例输入|样例输出|提示|模版代码|答案代码|测试用例|---)", p):
            break
        lines.append(p)

    result = "\n".join(lines).strip()
    return result if result else None


def _parse_test_cases(paragraphs: list) -> list:
    """解析测试用例段落，返回 [{"input": ..., "expected": ...}, ...]。"""
    cases = []
    current = {}
    mode = None  # "input" or "expected"

    for p in paragraphs:
        # 用例分隔符 --- 用例 N ---
        if p.startswith("---"):
            if current.get("input") and current.get("expected"):
                cases.append(current)
            current = {}
            mode = None
            continue

        # 输入：xxx
        m_in = re.match(r"输入[：:]\s*(.*)", p)
        if m_in:
            current["input"] = m_in.group(1).strip() + "\n"
            mode = "input"
            continue

        # 输出：xxx
        m_out = re.match(r"输出[：:]\s*(.*)", p)
        if m_out:
            current["expected"] = m_out.group(1).strip() + "\n"
            mode = "expected"
            continue

        # 多行续接
        if mode == "input" and "input" in current:
            current["input"] += p + "\n"
        elif mode == "expected" and "expected" in current:
            current["expected"] += p + "\n"

    if current.get("input") and current.get("expected"):
        cases.append(current)

    # 清理换行
    for c in cases:
        c["input"] = c["input"].strip()
        c["expected"] = c["expected"].strip()
    return cases


@app.route("/api/questions/upload-docx", methods=["POST"])
def upload_docx_question():
    """上传 Word 文档，解析后创建题目。"""
    if "file" not in request.files:
        return jsonify({"error": "缺少文件"}), 400
    file = request.files["file"]
    if file.filename == "":
        return jsonify({"error": "文件名为空"}), 400

    try:
        data = _parse_docx(file)
    except ValueError as e:
        return jsonify({"error": f"解析失败: {e}"}), 400
    except Exception as e:
        traceback.print_exc()
        _log("ERROR", "文档解析异常", str(e))
        return jsonify({"error": f"文档解析错误: {e}"}), 400

    # 填入 question 结构
    global questions
    new_id = max(q["id"] for q in questions) + 1 if questions else 1
    new_question = {
        "id": new_id,
        "title": data["title"],
        "difficulty": data["difficulty"],
        "category": data["category"],
        "description": data["description"],
        "input_format": data["input_format"],
        "output_format": data["output_format"],
        "sample_input": data["sample_input"],
        "sample_output": data["sample_output"],
        "hint": data.get("hint", ""),
        "template_code": data["template_code"],
        "template_code_cpp": STD_TEMPLATE_CPP,
        "test_cases": data["test_cases"],
    }
    questions.append(new_question)
    _save_questions(questions)

    return jsonify({"message": f"题目 #{new_id} 已添加", "id": new_id}), 201


@app.route("/api/questions/download-template")
def download_template():
    """下载 Word 题目模板。"""
    doc = Document()
    doc.add_heading("ZUWE-ACM 题目模板", level=1)
    doc.add_paragraph("标有 * 的为必填字段。测试用例中的 输入/输出 均为纯文本，如需多行则逐行续接即可。")
    doc.add_paragraph("")

    doc.add_heading("必填字段", level=2)
    doc.add_paragraph('标题*：两数之和')
    doc.add_paragraph('难度*：简单')
    doc.add_paragraph('分类*：基础算法-模拟')
    doc.add_paragraph('题目描述*：给定两个整数 a 和 b，计算它们的和。')
    doc.add_paragraph('输入格式*：一行两个整数 a 和 b，以空格分隔。')
    doc.add_paragraph('输出格式*：一行一个整数，表示 a 与 b 的和。')
    doc.add_paragraph('样例输入*：1 2')
    doc.add_paragraph('样例输出*：3')
    doc.add_paragraph('提示：注意输入范围，结果可能超过 32 位整数。')

    doc.add_paragraph("")
    doc.add_heading("测试用例", level=2)
    doc.add_paragraph("每个用例以 --- 分隔，输入和输出各占一行，冒号后跟内容：")
    doc.add_paragraph("")

    doc.add_paragraph("--- 用例 1 ---")
    doc.add_paragraph("输入：1 2")
    doc.add_paragraph("输出：3")

    doc.add_paragraph("--- 用例 2 ---")
    doc.add_paragraph("输入：10 20")
    doc.add_paragraph("输出：30")

    doc.add_paragraph("--- 用例 3 ---")
    doc.add_paragraph("输入：100 200")
    doc.add_paragraph("输出：300")

    doc.add_paragraph("")
    doc.add_heading("多行输入/输出的写法", level=2)
    doc.add_paragraph("如果测试用例的输入或输出需要多行，直接续行即可：")
    doc.add_paragraph("")
    doc.add_paragraph("--- 多行示例 ---")
    doc.add_paragraph("输入：3")
    doc.add_paragraph("1")
    doc.add_paragraph("2")
    doc.add_paragraph("3")
    doc.add_paragraph("输出：6")

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return Response(
        buf.read(),
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": "attachment; filename=ZUWE-ACM-Template.docx"},
    )


# ── AI 出题（DeepSeek）──

AI_API_KEY = os.environ.get("ANTHROPIC_AUTH_TOKEN", "")
AI_BASE_URL = os.environ.get("ANTHROPIC_BASE_URL", "https://api.deepseek.com/anthropic").rstrip("/")
AI_MODEL = os.environ.get("ANTHROPIC_MODEL", "DeepSeek-V4-flash")


@app.route("/api/ai/generate-question", methods=["POST"])
def ai_generate_question():
    """AI 根据描述自动生成题目并上传。"""
    data = request.get_json()
    description = data.get("description", "")
    api_key = data.get("api_key", "").strip() or AI_API_KEY or _settings.get("api_key", "").strip()

    if not api_key:
        return jsonify({"error": "请在设置中配置 API Key"}), 503
    if not description:
        return jsonify({"error": "请输入题目描述"}), 400

    system = """你是一个算法竞赛出题专家。用户会发一段文本给你，你需要自动识别它是什么：

1. 如果是一段完整的题目描述（可能包含输入输出格式、样例），则基于它生成完整的编程题
2. 如果是一个简单的想法或关键词，则自动扩展成一道合理的题目
3. 如果是 LeetCode / 牛客 / 其他 OJ 的题目链接或描述，则识别并重构为一道独立题目

请严格按照以下 JSON 格式输出，不要包含其他内容（不要用 markdown 代码块包裹）：

{
    "title": "题目标题",
    "difficulty": "简单 / 中等 / 困难",
    "category": "分类（如：基础算法-排序、数学-数论、数据结构-链表、模拟、字符串、贪心、动态规划、搜索、图论）",
    "description": "完整的题目描述",
    "input_format": "输入格式说明",
    "output_format": "输出格式说明",
    "sample_input": "样例输入（包含多行时用 \\n 表示换行）",
    "sample_output": "样例输出（包含多行时用 \\n 表示换行）",
    "hint": "提示信息（选填，没有就留空字符串）",
    "template_code": "Python 模板代码，使用 def main(): 格式，末尾有 if __name__ == '__main__': main()",
    "template_code_cpp": "C++ 模板代码，使用 #include <bits/stdc++.h> using namespace std; int main() { ... return 0; } 格式",
    "test_cases": [
        {"input": "测试输入1", "expected": "期望输出1"},
        {"input": "测试输入2", "expected": "期望输出2"},
        {"input": "测试输入3", "expected": "期望输出3"}
    ],
}

要求：
1. 至少生成 3 个测试用例，覆盖正常、边界和特殊情况
2. Python 代码用 def main(): 格式；C++ 代码用 int main() 格式
3. 难度要合理评估
4. 用中文填写所有文本字段，代码字段用英文"""

    try:
        # 延长心跳窗口，避免长时 AI 请求被 watchdog 杀死
        global _last_heartbeat
        _last_heartbeat = _time.time()
        resp = http_requests.post(
            f"{AI_BASE_URL}/messages",
            headers={
                "Content-Type": "application/json",
                "x-api-key": api_key,
                "anthropic-version": "2023-06-01",
            },
            json={
                "model": AI_MODEL,
                "system": system,
                "max_tokens": 4096,
                "messages": [{"role": "user", "content": description}],
            },
            timeout=120,
        )
        resp.raise_for_status()
        result = resp.json()

        ai_text = ""
        for block in result.get("content", []):
            if block.get("type") == "text":
                ai_text = block.get("text", "")
                break

        if not ai_text:
            return jsonify({"error": "AI 返回为空"}), 500

        clean = re.sub(r"^```(?:json)?\s*", "", ai_text.strip())
        clean = re.sub(r"\s*```$", "", clean)
        q_data = json.loads(clean)

        required = ["title", "difficulty", "category", "description",
                     "input_format", "output_format", "sample_input", "sample_output",
                     "template_code", "test_cases"]
        for field in required:
            if field not in q_data:
                return jsonify({"error": f"AI 生成的题目缺少字段: {field}"}), 500

        if not isinstance(q_data["test_cases"], list) or len(q_data["test_cases"]) == 0:
            return jsonify({"error": "AI 生成的测试用例无效"}), 500

        global questions
        new_id = _next_question_id()
        new_question = {
            "id": new_id,
            "title": q_data["title"],
            "difficulty": q_data["difficulty"],
            "category": q_data["category"],
            "description": q_data["description"],
            "input_format": q_data["input_format"],
            "output_format": q_data["output_format"],
            "sample_input": q_data["sample_input"],
            "sample_output": q_data["sample_output"],
            "hint": q_data.get("hint", ""),
            "template_code": q_data["template_code"],
            "template_code_cpp": q_data.get("template_code_cpp", ""),
            "test_cases": q_data["test_cases"],
        }
        questions.append(new_question)
        _save_questions(questions)

        return jsonify({
            "message": f"题目 #{new_id} 已通过 AI 创建",
            "id": new_id,
            "question": {
                "title": q_data["title"],
                "difficulty": q_data["difficulty"],
                "category": q_data["category"],
                "test_case_count": len(q_data["test_cases"]),
            },
        }), 201

    except json.JSONDecodeError:
        return jsonify({"error": "AI 返回格式异常，请重试。可能是 AI 输出不符合预期格式"}), 500
    except http_requests.exceptions.HTTPError as e:
        code = e.response.status_code if e.response is not None else 500
        if code == 401:
            return jsonify({"error": "API Key 无效或已过期，请在设置中重新配置"}), 503
        if code == 429:
            return jsonify({"error": "API 请求过于频繁，请稍后重试"}), 503
        return jsonify({"error": f"AI 服务返回错误 ({code})，请稍后重试"}), 503
    except http_requests.exceptions.Timeout:
        return jsonify({"error": "AI 服务响应超时，请检查网络后重试"}), 503
    except http_requests.exceptions.ConnectionError:
        return jsonify({"error": "无法连接到 AI 服务，请检查网络连接"}), 503
    except Exception as e:
        traceback.print_exc()
        _log("ERROR", "AI 生成题目异常", str(e))
        return jsonify({"error": f"生成失败: {type(e).__name__}"}), 500


@app.route("/api/ai/chat", methods=["POST"])
def ai_chat():
    """AI 对话（用于设置页面的 Key 验证和 AI 助教）。"""
    data = request.get_json()
    api_key = (data.get("api_key", "") or _settings.get("api_key", "") or "").strip()
    messages = data.get("messages", [])

    if not api_key:
        return jsonify({"error": "请在设置中配置 API Key"}), 503
    if not messages:
        return jsonify({"error": "消息不能为空"}), 400

    try:
        global _last_heartbeat
        _last_heartbeat = _time.time()
        resp = http_requests.post(
            f"{AI_BASE_URL}/messages",
            headers={
                "Content-Type": "application/json",
                "x-api-key": api_key,
                "anthropic-version": "2023-06-01",
            },
            json={
                "model": AI_MODEL,
                "max_tokens": 1024,
                "messages": messages,
            },
            timeout=30,
        )
        resp.raise_for_status()
        result = resp.json()

        reply = ""
        for block in result.get("content", []):
            if block.get("type") == "text":
                reply = block.get("text", "")
                break

        return jsonify({"reply": reply})
    except http_requests.exceptions.HTTPError as e:
        code = e.response.status_code if e.response is not None else 500
        if code == 401:
            return jsonify({"error": "API Key 无效或已过期"}), 503
        return jsonify({"error": f"AI 服务返回错误 ({code})"}), 503
    except http_requests.exceptions.Timeout:
        return jsonify({"error": "AI 服务响应超时"}), 503
    except http_requests.exceptions.ConnectionError:
        return jsonify({"error": "无法连接到 AI 服务，请检查网络连接"}), 503
    except Exception as e:
        _log("ERROR", "AI 对话异常", str(e))
        return jsonify({"error": f"对话失败: {type(e).__name__}"}), 500



@app.route("/api/compiler-status")
def compiler_status():
    """检查 C++ 编译器状态。"""
    import subprocess
    compiler = _sandbox_runner.CPP_COMPILER
    try:
        proc = subprocess.run(
            [compiler, "--version"],
            capture_output=True, text=True, timeout=5,
            creationflags=subprocess.CREATE_NO_WINDOW if sys.platform == "win32" else 0
        )
        if proc.returncode == 0:
            first_line = proc.stdout.strip().splitlines()[0] if proc.stdout.strip() else "OK"
            return jsonify({"found": True, "path": compiler, "version": first_line})
        return jsonify({"found": False, "path": compiler, "error": proc.stderr.strip()})
    except FileNotFoundError:
        return jsonify({"found": False, "path": compiler, "error": "未找到编译器，请安装 MinGW"})
    except Exception as e:
        return jsonify({"found": False, "path": compiler, "error": str(e)})


@app.route("/api/settings", methods=["GET", "POST"])
def settings_api():
    """获取/保存用户设置。"""
    global _settings, _custom_compiler
    if request.method == "POST":
        data = request.get_json() or {}
        _settings.update(data)
        _save_settings(_settings)
        _custom_compiler = _settings.get("compiler_path", "").strip()
        if _custom_compiler:
            _sandbox_runner.CPP_COMPILER = _custom_compiler
        _log("INFO", f"更新设置")
        return jsonify({"ok": True})
    # GET：返回设置，对 API Key 做掩码处理
    out = dict(_settings)
    if "api_key" in out and out["api_key"]:
        key = out["api_key"]
        out["api_key"] = key[:4] + "****" + key[-4:] if len(key) > 8 else "****"
    return jsonify(out)


# ── 代码保存/恢复（用户编写的代码持久化）──
SAVED_CODES_PATH = os.path.join(USER_DATA_DIR, "saved_codes.json")

def _load_saved_codes():
    try:
        with open(SAVED_CODES_PATH, "r", encoding="utf-8") as f:
            return json.load(f)
    except (FileNotFoundError, json.JSONDecodeError):
        return {}

def _save_saved_codes(data):
    try:
        with open(SAVED_CODES_PATH, "w", encoding="utf-8") as f:
            json.dump(data, f, ensure_ascii=False)
    except OSError:
        pass

@app.route("/api/code/save", methods=["POST"])
def save_code():
    """保存用户编写的代码到服务端。"""
    data = request.get_json()
    qid = str(data.get("question_id", ""))
    lang = data.get("language", "cpp")
    code = data.get("code", "")
    if not qid:
        return jsonify({"error": "缺少题目 ID"}), 400
    all_codes = _load_saved_codes()
    if qid not in all_codes:
        all_codes[qid] = {}
    all_codes[qid][lang] = code
    _save_saved_codes(all_codes)
    return jsonify({"ok": True})

@app.route("/api/code/load/<int:qid>")
def load_code(qid):
    """加载用户之前保存的某题代码。"""
    all_codes = _load_saved_codes()
    q = all_codes.get(str(qid), {})
    return jsonify(q)

@app.route("/api/code/load-all")
def load_all_codes():
    """加载所有用户保存的代码。"""
    return jsonify(_load_saved_codes())


@app.route("/api/heartbeat")
def heartbeat():
    """前端定时心跳，超时未收到则自动退出。"""
    global _last_heartbeat
    if not _quitting:
        _last_heartbeat = _time.time()
    return jsonify({"ok": True})


@app.route("/api/quit", methods=["POST", "GET"])
def quit_app():
    """退出整个程序。"""
    global _quitting
    _quitting = True
    _last_heartbeat = 0  # 让 watchdog 立即触发
    return jsonify({"ok": True})


if __name__ == "__main__":
    print(f"[√] 加载了 {len(questions)} 道题目")
    print(f"[*] 启动服务: http://127.0.0.1:5001")
    app.run(debug=True, host="0.0.0.0", port=5001)
